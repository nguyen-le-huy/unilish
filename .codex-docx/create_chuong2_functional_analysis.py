from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("/Users/nguyenlehuy/Downloads/DATN2026-D2_CNTT_CNPM_0191966_NguyenLeHuy_Chuong2_PhanTichChucNang.docx")
ARTIFACT = Path("/private/tmp/unilish-chuong2-functional/artifact.md")


def set_run_font(run, *, bold=None, italic=None, size=13, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_common(paragraph, *, first_line=True, before=0, after=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = 1.15
    fmt.first_line_indent = Cm(0.75) if first_line else None


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, bold=bold, size=12)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "808080")
        borders.append(tag)
    tbl_pr.append(borders)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, italic=True, size=11, color=(100, 100, 100))


def add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 3)
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_run_font(r, bold=True, size=13)
    return p


def add_body(doc, text, *, before=0, after=4):
    p = doc.add_paragraph(style="Body Text")
    set_paragraph_common(p, before=before, after=after)
    r = p.add_run(text)
    set_run_font(r, size=13)
    return p


def add_label_paragraph(doc, label, text):
    p = doc.add_paragraph(style="Body Text")
    set_paragraph_common(p, before=2, after=3)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True, size=13)
    r2 = p.add_run(text)
    set_run_font(r2, size=13)
    return p


def add_bullet(doc, text, *, level=0):
    p = doc.add_paragraph(style="List Paragraph")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(1.25 + level * 0.5)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    r = p.add_run("- " + text)
    set_run_font(r, size=13)
    return p


def configure_document(doc):
    sec = doc.sections[0]
    sec.start_type = WD_SECTION.NEW_PAGE
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.75)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(1.9)
    sec.bottom_margin = Cm(1.9)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.25)

    for style_name in ["Normal", "Body Text", "List Paragraph", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        style.font.size = Pt(13)
        if style_name.startswith("Heading"):
            style.font.bold = True
        pf = style.paragraph_format
        pf.line_spacing = 1.15

    footer = sec.footer
    for p in footer.paragraphs:
        p.clear()
    line = footer.paragraphs[0]
    add_bottom_border(line)
    table = footer.add_table(rows=1, cols=2, width=Inches(6.4))
    table.autofit = False
    table.columns[0].width = Inches(4.6)
    table.columns[1].width = Inches(1.8)
    left = table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = left.add_run("Nguyễn Lê Huy - 0191966 - 66PM")
    set_run_font(r, italic=True, size=11, color=(100, 100, 100))
    add_page_number(table.cell(0, 1).paragraphs[0])


def build():
    doc = Document()
    configure_document(doc)

    add_heading(doc, "CHƯƠNG 2: PHÂN TÍCH, THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "1. Phân tích chức năng", 2)

    add_body(
        doc,
        "Phần phân tích chức năng nhằm xác định các nhóm nghiệp vụ chính mà hệ thống Unilish cần đáp ứng. "
        "Hệ thống được xây dựng cho hai nhóm người dùng chính là người học và quản trị viên, đồng thời có sự tham gia của các dịch vụ hỗ trợ như xác thực, lưu trữ dữ liệu, xử lý giọng nói và trí tuệ nhân tạo. "
        "Các chức năng được phân tích theo hướng phục vụ trực tiếp cho quá trình học ngoại ngữ trực tuyến: đánh giá trình độ, đề xuất lộ trình học, học theo bài, luyện tập kỹ năng và quản lý nội dung học tập.",
    )

    add_heading(doc, "1.1. Tác nhân sử dụng hệ thống", 2)
    add_body(
        doc,
        "Trong phạm vi đề tài, hệ thống gồm các tác nhân chính sau:",
    )
    actors = [
        ("Khách truy cập", "Xem thông tin giới thiệu hệ thống, đăng ký tài khoản hoặc đăng nhập để bắt đầu học."),
        ("Người học", "Sử dụng website để thiết lập mục tiêu, làm bài kiểm tra đầu vào, học khóa học, luyện kỹ năng và theo dõi tiến độ cá nhân."),
        ("Quản trị viên", "Quản lý người dùng, ngôn ngữ, mục tiêu học, khóa học, bài học, câu hỏi, bài kiểm tra đầu vào và nội dung luyện thi IELTS."),
        ("Content creator", "Tạo, chỉnh sửa, kiểm tra và xuất bản nội dung học tập theo quyền được phân công."),
        ("Hệ thống AI và dịch vụ ngoài", "Hỗ trợ sinh nội dung, gợi ý học tập, xử lý giọng nói, lưu trữ media và cung cấp dữ liệu cho các chức năng luyện nghe, nói."),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    widths = [Cm(4.2), Cm(12.0)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    set_cell_text(table.cell(0, 0), "Tác nhân", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), "Vai trò trong hệ thống", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(table.cell(0, 0), "D9EAF7")
    set_cell_shading(table.cell(0, 1), "D9EAF7")
    for actor, role in actors:
        row = table.add_row()
        set_cell_text(row.cells[0], actor, bold=True)
        set_cell_text(row.cells[1], role)

    add_heading(doc, "1.2. Nhóm chức năng dành cho người học", 2)
    add_body(
        doc,
        "Người học là tác nhân trung tâm của hệ thống. Các chức năng dành cho người học được tổ chức theo luồng sử dụng thực tế, bắt đầu từ tạo tài khoản, thiết lập hồ sơ học tập, làm bài đánh giá, nhận khóa học phù hợp và tiếp tục học theo lộ trình đã được hệ thống ghi nhận.",
    )
    add_label_paragraph(doc, "a. Đăng ký, đăng nhập và quản lý tài khoản: ", "Người học có thể đăng ký tài khoản bằng email, xác thực OTP, đăng nhập bằng email/mật khẩu hoặc Google OAuth. Sau khi đăng nhập, hệ thống cấp phiên truy cập để bảo vệ các trang cần xác thực. Người học có thể xem và cập nhật thông tin cá nhân, đồng thời hệ thống lưu trạng thái học tập gắn với tài khoản.")
    add_label_paragraph(doc, "b. Onboarding và thiết lập mục tiêu học tập: ", "Khi bắt đầu sử dụng, người học lựa chọn ngôn ngữ muốn học, mục tiêu học tập, trình độ hiện tại và trình độ mong muốn. Các thông tin này là cơ sở để hệ thống cá nhân hóa đề xuất khóa học và sắp xếp nội dung phù hợp.")
    add_label_paragraph(doc, "c. Kiểm tra trình độ đầu vào: ", "Người học thực hiện placement test để hệ thống đánh giá năng lực ban đầu. Kết quả bài kiểm tra được lưu lại, sử dụng để xác định cấp độ học, gợi ý khóa học và hỗ trợ theo dõi sự tiến bộ trong quá trình học.")
    add_label_paragraph(doc, "d. Nhận gợi ý khóa học: ", "Dựa trên mục tiêu, ngôn ngữ, trình độ và kết quả kiểm tra đầu vào, hệ thống đề xuất các khóa học phù hợp. Người học có thể xem danh sách khóa học, xem thông tin tổng quan và đăng ký học khóa học mong muốn.")
    add_label_paragraph(doc, "e. Học theo lộ trình khóa học: ", "Sau khi đăng ký khóa học, người học xem roadmap gồm các unit và lesson được sắp xếp theo thứ tự. Mỗi bài học có nội dung học tập, bài luyện tập, trạng thái hoàn thành và điểm số. Hệ thống hỗ trợ lưu checkpoint để người học có thể tiếp tục từ vị trí đang học.")
    add_label_paragraph(doc, "f. Làm bài tập và theo dõi kết quả: ", "Người học làm bài tập trong từng lesson, nộp bài, nhận kết quả chấm điểm và phản hồi. Hệ thống lưu số lần làm bài, điểm tốt nhất, trạng thái bài học và tiến độ tổng thể của khóa học.")
    add_label_paragraph(doc, "g. Luyện kỹ năng nghe, nói và phát âm: ", "Hệ thống cung cấp các hoạt động luyện AI Voice, shadowing và dictation. Người học có thể luyện hội thoại theo chủ đề, nghe và lặp lại nội dung, luyện chép chính tả và nhận hỗ trợ từ các dịch vụ xử lý giọng nói.")
    add_label_paragraph(doc, "h. Luyện IELTS theo từng kỹ năng: ", "Người học truy cập khu vực IELTS Practice để luyện Listening, Reading, Writing và Speaking. Hệ thống lưu attempt, kết quả làm bài và thông tin phục vụ đánh giá quá trình luyện tập.")

    add_heading(doc, "1.3. Nhóm chức năng dành cho quản trị viên", 2)
    add_body(
        doc,
        "Quản trị viên và người tạo nội dung sử dụng Admin CMS để quản lý dữ liệu học tập. Nhóm chức năng này giúp hệ thống có thể cập nhật nội dung mà không cần can thiệp trực tiếp vào mã nguồn, đồng thời đảm bảo dữ liệu học tập được tổ chức nhất quán.",
    )
    admin_funcs = [
        "Quản lý ngôn ngữ học, mã ngôn ngữ, cấu hình giọng đọc và các thông tin phục vụ nội dung đa ngôn ngữ.",
        "Quản lý mục tiêu học tập và chiến lược học, làm cơ sở cho quá trình gợi ý khóa học.",
        "Quản lý khóa học, unit và lesson; tạo cấu trúc chương trình học theo từng cấp độ và mục tiêu.",
        "Quản lý ngân hàng câu hỏi, bao gồm tạo mới, chỉnh sửa, duyệt, publish hoặc archive câu hỏi.",
        "Quản lý placement test để phục vụ đánh giá trình độ đầu vào cho người học.",
        "Quản lý đề luyện IELTS theo từng kỹ năng, phiên bản và trạng thái xuất bản.",
        "Quản lý người dùng, kiểm tra thông tin tài khoản và hỗ trợ xử lý các trường hợp cần can thiệp quản trị.",
        "Quản lý media học tập như âm thanh, hình ảnh, transcript hoặc nội dung liên quan đến bài học.",
    ]
    for item in admin_funcs:
        add_bullet(doc, item)

    add_heading(doc, "1.4. Nhóm chức năng hệ thống và dịch vụ hỗ trợ", 2)
    add_body(
        doc,
        "Ngoài các chức năng trực tiếp hiển thị trên giao diện, Unilish cần các chức năng nền để đảm bảo hệ thống hoạt động ổn định, an toàn và có khả năng mở rộng. Các chức năng này chủ yếu nằm ở backend API và các dịch vụ tích hợp.",
    )
    add_label_paragraph(doc, "a. Xác thực và phân quyền: ", "Backend kiểm tra thông tin đăng nhập, phát hành token, xác định vai trò người dùng và bảo vệ các API yêu cầu quyền truy cập. Các chức năng quản trị chỉ cho phép người dùng có quyền phù hợp thao tác.")
    add_label_paragraph(doc, "b. Lưu trữ và truy xuất dữ liệu: ", "Hệ thống lưu dữ liệu người dùng, khóa học, bài học, câu hỏi, tiến độ, attempt và kết quả học tập trong cơ sở dữ liệu. Các API được thiết kế để client và admin có thể truy xuất dữ liệu theo từng nghiệp vụ.")
    add_label_paragraph(doc, "c. Xử lý nội dung và trí tuệ nhân tạo: ", "Các dịch vụ AI hỗ trợ gợi ý học tập, xử lý giọng nói, tạo phản hồi hoặc phục vụ các tính năng luyện nghe nói. Việc tích hợp AI giúp hệ thống tăng tính tương tác thay vì chỉ cung cấp nội dung tĩnh.")
    add_label_paragraph(doc, "d. Quản lý tệp và media: ", "Hệ thống cần lưu trữ âm thanh, hình ảnh và các tài nguyên phục vụ lesson, shadowing, dictation hoặc IELTS Practice. Các tệp này được liên kết với nội dung học tập để người học truy cập trong quá trình sử dụng.")
    add_label_paragraph(doc, "e. Theo dõi trạng thái và xử lý lỗi: ", "Backend ghi log, chuẩn hóa phản hồi lỗi và kiểm soát trạng thái xử lý của các nghiệp vụ quan trọng. Điều này giúp quá trình vận hành, kiểm thử và bảo trì hệ thống thuận lợi hơn.")

    add_heading(doc, "1.5. Tổng hợp yêu cầu chức năng", 2)
    summary = [
        ("F01", "Quản lý tài khoản", "Đăng ký, đăng nhập, OTP, Google OAuth, cập nhật hồ sơ và phân quyền truy cập."),
        ("F02", "Thiết lập hồ sơ học", "Chọn ngôn ngữ, mục tiêu, trình độ hiện tại và trình độ mong muốn."),
        ("F03", "Kiểm tra trình độ", "Làm placement test, chấm kết quả, lưu attempt và xác định cấp độ phù hợp."),
        ("F04", "Gợi ý khóa học", "Đề xuất khóa học dựa trên mục tiêu, trình độ và dữ liệu học tập của người học."),
        ("F05", "Học theo lộ trình", "Xem roadmap, mở lesson, học nội dung, làm bài tập, lưu checkpoint và cập nhật tiến độ."),
        ("F06", "Luyện kỹ năng", "Hỗ trợ AI Voice, shadowing, dictation và IELTS Practice theo từng kỹ năng."),
        ("F07", "Quản trị nội dung", "Quản lý ngôn ngữ, mục tiêu, khóa học, unit, lesson, câu hỏi và đề kiểm tra."),
        ("F08", "Quản lý người dùng", "Theo dõi thông tin người dùng, vai trò, trạng thái tài khoản và dữ liệu học tập liên quan."),
        ("F09", "Tích hợp dịch vụ hỗ trợ", "Kết nối dịch vụ AI, speech, lưu trữ media, logging và các thành phần backend cần thiết."),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)
    for i, title in enumerate(["Mã", "Nhóm chức năng", "Mô tả yêu cầu"]):
        set_cell_text(table.cell(0, i), title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.cell(0, i), "D9EAF7")
    for code, group, desc in summary:
        row = table.add_row()
        set_cell_text(row.cells[0], code, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], group)
        set_cell_text(row.cells[2], desc)

    add_body(
        doc,
        "Như vậy, phần phân tích chức năng cho thấy Unilish không chỉ là một website hiển thị bài học, mà là một hệ thống học ngoại ngữ có đầy đủ luồng nghiệp vụ từ đánh giá đầu vào, cá nhân hóa lộ trình, học và luyện tập đến quản trị nội dung. Các chức năng này là cơ sở để tiếp tục phân tích hệ thống bằng biểu đồ use case, biểu đồ tuần tự, biểu đồ lớp và thiết kế cơ sở dữ liệu ở các phần tiếp theo của chương.",
        before=8,
    )

    doc.save(OUT)

    ARTIFACT.parent.mkdir(parents=True, exist_ok=True)
    ARTIFACT.write_text(
        "\n".join(
            [
                "# Artifact contract",
                f"Reference: /Users/nguyenlehuy/Downloads/DATN2026-D2_CNTT_CNPM_0191966_NguyenLeHuy_BaoCao.docx",
                f"Output: {OUT}",
                "Page: A4 portrait, margins based on source report: left 2.75cm, right 2.0cm, top 1.9cm, bottom 1.9cm.",
                "Typography: Times New Roman 13pt for headings/body/list text; footer 11pt italic.",
                "Headings: Heading 1 centered bold 13pt; Heading 2 bold 13pt justified.",
                "Body: justified, 1.15 line spacing, first-line indent 0.75cm.",
                "Footer: horizontal rule, left student identity, right PAGE field.",
                "Content: only Chapter 2 section 1 functional analysis; section 2 system analysis intentionally omitted.",
            ]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    build()
