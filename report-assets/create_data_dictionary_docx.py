from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = "/Users/nguyenlehuy/Downloads/Unilish_BangMoTaDuLieu_TuERD.docx"


tables = [
    ("language", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của ngôn ngữ."),
        ("name", "VARCHAR", "255", "Not null", "Tên ngôn ngữ hiển thị trong hệ thống."),
        ("nativeName", "VARCHAR", "255", "Not null", "Tên bản địa của ngôn ngữ."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của ngôn ngữ."),
        ("flagIconUrl", "VARCHAR", "255", "Nullable", "Đường dẫn biểu tượng cờ của ngôn ngữ."),
        ("greeting", "VARCHAR", "255", "Nullable", "Câu chào mẫu theo ngôn ngữ."),
        ("greetingSoundUrl", "VARCHAR", "255", "Nullable", "Đường dẫn âm thanh câu chào."),
        ("learningGoalsId", "INTEGER", "10", "FK", "Mã mục tiêu học tập liên kết."),
    ]),
    ("learningGoal", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của mục tiêu học tập."),
        ("title", "VARCHAR", "255", "Not null", "Tên mục tiêu học tập."),
        ("systemPrompt", "VARCHAR", "255", "Nullable", "Prompt hệ thống phục vụ gợi ý/luyện tập bằng AI."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của mục tiêu."),
        ("supportedLanguages", "INTEGER", "10", "Nullable", "Danh sách hoặc mã ngôn ngữ được hỗ trợ."),
        ("iconUrl", "VARCHAR", "255", "Nullable", "Đường dẫn biểu tượng mục tiêu."),
        ("description", "VARCHAR", "255", "Nullable", "Mô tả chi tiết mục tiêu học tập."),
    ]),
    ("user", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của người dùng."),
        ("name", "VARCHAR", "255", "Not null", "Tên đăng nhập hoặc tên hiển thị."),
        ("password", "VARCHAR", "255", "Nullable", "Mật khẩu đã mã hóa, có thể rỗng nếu đăng nhập Google."),
        ("isVerified", "BIT", "1", "Not null", "Trạng thái xác minh tài khoản."),
        ("role", "VARCHAR", "255", "Not null", "Vai trò người dùng như Learner hoặc Admin."),
        ("fullName", "VARCHAR", "255", "Nullable", "Họ tên đầy đủ của người dùng."),
        ("avatarUrl", "VARCHAR", "255", "Nullable", "Đường dẫn ảnh đại diện."),
        ("lastActiveCourseId", "VARCHAR", "255", "Nullable", "Khóa học được truy cập gần nhất."),
        ("learningGoalID", "INTEGER", "10", "FK, Nullable", "Mục tiêu học tập hiện tại của người dùng."),
        ("currentLevel", "VARCHAR", "255", "Nullable", "Trình độ hiện tại của người dùng."),
        ("googleId", "VARCHAR", "255", "Nullable", "Mã định danh Google OAuth."),
        ("placementTestScore", "INTEGER", "10", "Nullable", "Điểm bài kiểm tra đầu vào."),
        ("learningLanguageId", "INTEGER", "10", "FK, Nullable", "Ngôn ngữ người dùng đang học."),
    ]),
    ("course", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của khóa học."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề khóa học."),
        ("slug", "VARCHAR", "255", "Not null", "Đường dẫn thân thiện của khóa học."),
        ("name", "VARCHAR", "255", "Not null", "Tên khóa học."),
        ("level", "VARCHAR", "255", "Nullable", "Cấp độ của khóa học."),
        ("languageId", "INTEGER", "10", "FK", "Ngôn ngữ của khóa học."),
        ("learningGoalId", "INTEGER", "10", "FK", "Mục tiêu học tập mà khóa học phục vụ."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của khóa học."),
    ]),
    ("unit", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của học phần."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề học phần."),
        ("orderIndex", "INTEGER", "10", "Not null", "Thứ tự hiển thị của học phần trong khóa học."),
        ("courseId", "INTEGER", "10", "FK", "Khóa học chứa học phần."),
    ]),
    ("lesson", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của bài học."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề bài học."),
        ("type", "VARCHAR", "255", "Not null", "Loại bài học, ví dụ vocabulary, grammar, listening."),
        ("content", "VARCHAR", "255", "Nullable", "Nội dung bài học."),
        ("orderIndex", "INTEGER", "10", "Not null", "Thứ tự bài học trong học phần."),
        ("unitId", "INTEGER", "10", "FK", "Học phần chứa bài học."),
    ]),
    ("courseEnrollment", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của bản ghi ghi danh."),
        ("status", "VARCHAR", "255", "Not null", "Trạng thái học khóa học."),
        ("timeSpentSeconds", "INTEGER", "10", "Not null", "Tổng thời gian học tính bằng giây."),
        ("completedLessonCount", "INTEGER", "10", "Not null", "Số bài học đã hoàn thành."),
        ("userId", "INTEGER", "10", "FK", "Người học tham gia khóa học."),
        ("courseId", "INTEGER", "10", "FK", "Khóa học được ghi danh."),
        ("lastLessonId", "INTEGER", "10", "FK, Nullable", "Bài học truy cập gần nhất."),
    ]),
    ("userLessonProgress", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của tiến độ bài học."),
        ("name", "VARCHAR", "255", "Nullable", "Tên hoặc nhãn tiến độ theo dữ liệu ERD."),
        ("password", "VARCHAR", "255", "Nullable", "Trường dữ liệu theo ERD hiện tại."),
        ("userId", "INTEGER", "10", "FK", "Người học sở hữu tiến độ."),
        ("lessonId", "INTEGER", "10", "FK", "Bài học được theo dõi tiến độ."),
    ]),
    ("placementTest", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của bài kiểm tra đầu vào."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề bài kiểm tra."),
        ("name", "VARCHAR", "255", "Not null", "Tên bài kiểm tra."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của bài kiểm tra."),
        ("standard", "VARCHAR", "255", "Nullable", "Chuẩn đánh giá hoặc thang điểm."),
        ("languageId", "INTEGER", "10", "FK", "Ngôn ngữ áp dụng cho bài kiểm tra."),
    ]),
    ("question", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của câu hỏi."),
        ("skill", "VARCHAR", "255", "Not null", "Kỹ năng liên quan đến câu hỏi."),
        ("type", "VARCHAR", "255", "Not null", "Loại câu hỏi."),
        ("content", "VARCHAR", "255", "Not null", "Nội dung câu hỏi."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của câu hỏi."),
        ("placementTestId", "INTEGER", "10", "FK", "Bài kiểm tra đầu vào chứa câu hỏi."),
    ]),
    ("placementTestAttempt", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của lượt làm bài đầu vào."),
        ("status", "VARCHAR", "255", "Not null", "Trạng thái lượt làm bài."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của lượt làm."),
        ("startedAt", "DATE", "—", "Not null", "Thời điểm bắt đầu làm bài."),
        ("expiredAt", "DATE", "—", "Nullable", "Thời điểm hết hạn làm bài."),
        ("scoring", "INTEGER", "10", "Nullable", "Điểm số của lượt làm bài."),
        ("placementTestId", "INTEGER", "10", "FK", "Bài kiểm tra được thực hiện."),
        ("userId", "INTEGER", "10", "FK", "Người dùng thực hiện bài kiểm tra."),
    ]),
    ("shadowingVideo", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của video luyện nói/nối đuôi."),
        ("videoID", "VARCHAR", "255", "Not null", "Mã video YouTube hoặc mã video hệ thống."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề video."),
        ("thumbnailUrl", "VARCHAR", "255", "Nullable", "Đường dẫn ảnh thumbnail."),
        ("isActive", "BIT", "1", "Not null", "Trạng thái hoạt động của video."),
        ("durationSeconds", "INTEGER", "10", "Nullable", "Thời lượng video tính bằng giây."),
        ("cues", "TINYINT", "3", "Nullable", "Dữ liệu cue/đoạn luyện tập theo ERD."),
        ("userId", "INTEGER", "10", "FK", "Người dùng tạo hoặc luyện tập video."),
    ]),
    ("examTest", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của đề thi/luyện tập IELTS."),
        ("format", "VARCHAR", "255", "Not null", "Định dạng đề thi."),
        ("title", "VARCHAR", "255", "Not null", "Tiêu đề đề thi."),
        ("skill", "VARCHAR", "255", "Not null", "Kỹ năng IELTS."),
        ("status", "VARCHAR", "255", "Not null", "Trạng thái đề thi."),
        ("content", "VARCHAR", "255", "Nullable", "Nội dung/cấu trúc đề thi."),
        ("languageId", "INTEGER", "10", "FK", "Ngôn ngữ của đề thi."),
    ]),
    ("ieltsPracticeAttempt", [
        ("id", "INTEGER", "10", "Primary key, auto-increment", "Khóa chính của lượt luyện đề IELTS."),
        ("format", "VARCHAR", "255", "Not null", "Định dạng bài luyện."),
        ("questionType", "VARCHAR", "255", "Not null", "Loại câu hỏi."),
        ("skill", "VARCHAR", "255", "Not null", "Kỹ năng IELTS được luyện tập."),
        ("status", "VARCHAR", "255", "Not null", "Trạng thái lượt luyện tập."),
        ("result", "INTEGER", "10", "Nullable", "Kết quả hoặc điểm số lượt luyện tập."),
        ("userId", "INTEGER", "10", "FK", "Người dùng thực hiện luyện tập."),
        ("examTestId", "INTEGER", "10", "FK", "Đề IELTS được luyện tập."),
    ]),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def set_font(paragraph, size=10.5, bold=False, italic=False):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.25)
section.footer_distance = Cm(1.25)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
styles["Normal"].font.size = Pt(13)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(12)
r = title.add_run("BẢNG MÔ TẢ CƠ SỞ DỮ LIỆU HỆ THỐNG UNILISH")
r.bold = True
r.font.name = "Times New Roman"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
r.font.size = Pt(14)

intro = doc.add_paragraph(
    "Phần này mô tả cấu trúc dữ liệu chính của hệ thống Unilish dựa trên biểu đồ ERD, "
    "bao gồm tên trường, kiểu dữ liệu, độ dài, ràng buộc và ý nghĩa sử dụng của từng trường."
)
intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
intro.paragraph_format.line_spacing = 1.3
set_font(intro, size=13)

headers = ["FieldName", "Data Type", "Data Length", "Constraint", "Description"]
widths = [3.0, 2.2, 1.8, 3.3, 6.1]

for index, (name, rows) in enumerate(tables, start=1):
    heading = doc.add_paragraph()
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(f"3.3.{index}. Bảng {name}")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_cell_margins(table)
    set_table_width(table, widths)

    hdr_cells = table.rows[0].cells
    repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for col, text in enumerate(headers):
        hdr_cells[col].text = text
        set_cell_shading(hdr_cells[col], "D9D9D9")
        hdr_cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr_cells[col].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p, size=10.5, bold=True)

    for row in rows:
        docx_row = table.add_row()
        prevent_row_split(docx_row)
        cells = docx_row.cells
        for col, text in enumerate(row):
            cells[col].text = text
            cells[col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[col].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if col in (1, 2):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(p, size=10.5)
        set_table_width(table, widths)

    add_caption(doc, f"Bảng 3.{index}. Bảng {name}")

doc.save(OUTPUT)
print(OUTPUT)
