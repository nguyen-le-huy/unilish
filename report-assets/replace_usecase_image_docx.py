from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOCX = Path("/Users/nguyenlehuy/Downloads/DATN2026-D2_CNTT_CNPM_0191966_NguyenLeHuy_BaoCao_SuaUseCaseTrang32.docx")
IMG = Path("/Users/nguyenlehuy/Downloads/unilish/report-assets/unilish-usecase-simple.png")


def set_run_font(run, size=13, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.italic = italic


def body_text(element):
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def main():
    doc = Document(DOCX)
    target_paragraph = None
    caption_paragraph = None

    for i, p in enumerate(doc.paragraphs):
        if p._element.findall(".//" + qn("w:drawing")):
            # The target image is the first drawing immediately after the Use case heading.
            prev_text = doc.paragraphs[i - 1].text.strip() if i > 0 else ""
            next_text = doc.paragraphs[i + 1].text.strip() if i + 1 < len(doc.paragraphs) else ""
            if prev_text == "Use case" or "Use case Diagram tổng quát" in next_text:
                target_paragraph = p
                if i + 1 < len(doc.paragraphs):
                    caption_paragraph = doc.paragraphs[i + 1]
                break

    if target_paragraph is None:
        raise RuntimeError("Không tìm thấy ảnh Use Case cần thay.")

    target_paragraph.clear()
    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = target_paragraph.add_run()
    run.add_picture(str(IMG), width=Cm(14.0))

    if caption_paragraph is not None:
        caption_paragraph.clear()
        caption_paragraph.style = "Caption"
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = caption_paragraph.add_run("Hình 2.1. Use Case Diagram tổng quát hệ thống Unilish")
        set_run_font(r, 13, italic=True)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
