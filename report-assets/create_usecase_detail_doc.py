from html import escape
from pathlib import Path
import subprocess

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE = Path("/Users/nguyenlehuy/Downloads/unilish/report-assets/usecase-detail")
BASE.mkdir(parents=True, exist_ok=True)
OUT = Path("/Users/nguyenlehuy/Downloads/Unilish_UseCase_Detail_SauTrang47.docx")
NODE = "/Users/nguyenlehuy/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/bin/node"
NODE_PATH = "/Users/nguyenlehuy/.cache/codex-runtimes/codex-primary-runtime/dependencies/node/node_modules"


USECASES = [
    ("Đăng ký", "Learner", ["Nhập thông tin đăng ký", "Xác thực OTP", "Tạo tài khoản", "Đăng ký bằng Google OAuth"]),
    ("Đăng nhập", "Learner, Admin", ["Đăng nhập bằng email và password", "Đăng nhập bằng Google OAuth", "Xác thực OTP", "Tạo phiên đăng nhập"]),
    ("Làm bài kiểm tra đầu vào", "Learner", ["Trả lời câu hỏi", "Nộp bài", "Xem kết quả", "Nhận lộ trình học cá nhân hóa"]),
    ("Luyện kỹ năng giao tiếp với AI", "Learner", ["Chọn chủ đề giao tiếp", "Gửi câu trả lời", "Nhận phản hồi từ AI", "Xem gợi ý cải thiện"]),
    ("Học theo khóa học được đề xuất", "Learner", ["Xem danh sách nội dung", "Học bài học", "Làm bài luyện tập", "Theo dõi tiến độ", "Nhận đề xuất từ AI"]),
    ("Học với video YouTube", "Learner", ["Xem video bài học", "Nghe chép chính tả", "Nói đuổi", "Cập nhật tiến độ"]),
    ("Quản lý profile", "Learner", ["Cập nhật thông tin cá nhân", "Chọn ngôn ngữ học", "Chọn mục tiêu học tập", "Chọn trình độ hiện tại"]),
    ("Luyện đề IELTS", "Learner", ["Chọn kỹ năng IELTS", "Làm bài luyện tập", "Nộp bài", "Xem kết quả luyện đề"]),
    ("Quản lý ngôn ngữ", "Admin", ["Tạo ngôn ngữ", "Cập nhật ngôn ngữ", "Xóa hoặc ẩn ngôn ngữ"]),
    ("Quản lý mục tiêu", "Admin", ["Tạo mục tiêu", "Cập nhật mục tiêu", "Xóa hoặc ẩn mục tiêu"]),
    ("Quản lý nội dung học", "Admin", ["Quản lý Course Series", "Quản lý Course", "Quản lý Unit", "Quản lý Lesson", "Xuất bản nội dung"]),
    ("Quản lý user", "Admin", ["Xem danh sách User", "Xem thông tin chi tiết User", "Quản lý quyền", "Khóa hoặc mở khóa tài khoản"]),
    ("Quản lý ngân hàng câu hỏi", "Admin", ["Tạo câu hỏi", "Cập nhật câu hỏi", "Xóa hoặc lưu trữ câu hỏi", "Duyệt câu hỏi"]),
    ("Quản lý bài kiểm tra đầu vào", "Admin", ["Tạo bài kiểm tra đầu vào", "Cập nhật bài kiểm tra đầu vào", "Xóa hoặc lưu trữ bài kiểm tra", "Cấu hình thang điểm"]),
    ("Quản lý đề IELTS", "Admin", ["Tạo đề", "Cập nhật đề", "Xóa hoặc lưu trữ đề", "Xuất bản đề IELTS"]),
]


def wrap_words(s, width=24):
    words = s.split()
    lines, line = [], ""
    for word in words:
        candidate = word if not line else f"{line} {word}"
        if len(candidate) <= width:
            line = candidate
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return lines


def svg_text(x, y, s, size=18, weight="400", anchor="middle", italic=False):
    style = "font-style:italic;" if italic else ""
    return (
        f'<text x="{x}" y="{y}" text-anchor="{anchor}" font-family="Times New Roman, serif" '
        f'font-size="{size}" font-weight="{weight}" fill="#000" style="{style}">{escape(s)}</text>'
    )


def multiline(x, y, s, size=17, width=24):
    lines = wrap_words(s, width)
    start = y - (len(lines) - 1) * size * 0.55
    spans = []
    for i, line in enumerate(lines):
        spans.append(f'<tspan x="{x}" y="{start + i * size * 1.15}">{escape(line)}</tspan>')
    return f'<text text-anchor="middle" font-family="Times New Roman, serif" font-size="{size}" fill="#000">{"".join(spans)}</text>'


def actor_svg(x, y, label):
    label_lines = label.split(", ")
    labels = []
    for i, line in enumerate(label_lines):
        labels.append(svg_text(x, y + 105 + i * 22, line, 16))
    return (
        f'<circle cx="{x}" cy="{y-42}" r="17" fill="none" stroke="#000" stroke-width="1.4"/>'
        f'<line x1="{x}" y1="{y-25}" x2="{x}" y2="{y+30}" stroke="#000" stroke-width="1.4"/>'
        f'<line x1="{x-34}" y1="{y-4}" x2="{x+34}" y2="{y-4}" stroke="#000" stroke-width="1.4"/>'
        f'<line x1="{x}" y1="{y+30}" x2="{x-32}" y2="{y+78}" stroke="#000" stroke-width="1.4"/>'
        f'<line x1="{x}" y1="{y+30}" x2="{x+32}" y2="{y+78}" stroke="#000" stroke-width="1.4"/>'
        + "".join(labels)
    )


def oval_svg(x, y, label):
    return (
        f'<ellipse cx="{x}" cy="{y}" rx="142" ry="33" fill="#fff" stroke="#000" stroke-width="1.35"/>'
        + multiline(x, y + 6, label, 16, 25)
    )


def line_svg(x1, y1, x2, y2):
    return f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="#000" stroke-width="1.25"/>'


def make_svg(index, title, actor, items):
    w, h = 920, 520
    if len(items) >= 5:
        h = 620
    top = 130
    step = 88
    ys = [top + i * step for i in range(len(items))]
    actor_y = (ys[0] + ys[-1]) / 2

    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{w}" height="{h}" viewBox="0 0 {w} {h}">',
        '<rect width="100%" height="100%" fill="#fff"/>',
        f'<rect x="145" y="1" width="{w-146}" height="{h-2}" fill="none" stroke="#000" stroke-width="1.5"/>',
        multiline(545, 33, f"Usecase {title}", 18, 28),
        actor_svg(55, actor_y, actor),
    ]
    for label, y in zip(items, ys):
        parts.append(oval_svg(535, y, label))
        parts.append(line_svg(90, actor_y, 393, y))
    parts.append("</svg>")
    svg_path = BASE / f"usecase-{index:02d}.svg"
    svg_path.write_text("\n".join(parts), encoding="utf-8")
    return svg_path


def convert_svg_to_png(svg_path):
    png_path = svg_path.with_suffix(".png")
    script = f"""
const sharp = require('sharp');
(async () => {{
  await sharp({str(svg_path)!r})
    .resize({{ width: 1700 }})
    .png()
    .toFile({str(png_path)!r});
}})();
"""
    subprocess.run(
        [NODE, "-e", script],
        check=True,
        env={"NODE_PATH": NODE_PATH},
        cwd=str(BASE),
    )
    return png_path


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_top_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)


def add_heading(doc, text_value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text_value)
    set_font(r, 13, bold=True)
    return p


def add_caption(doc, text_value):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text_value)
    set_font(r, 13, italic=True)


def build_doc(image_paths):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)

    styles = doc.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Caption"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)

    top = doc.add_paragraph()
    add_top_border(top)
    top.paragraph_format.space_after = Pt(8)

    h1 = doc.add_paragraph()
    h1.paragraph_format.space_after = Pt(3)
    r = h1.add_run("2.2. Use case Diagram chi tiết")
    set_font(r, 13, bold=True)

    intro = doc.add_paragraph()
    intro.paragraph_format.first_line_indent = Cm(1.0)
    intro.paragraph_format.line_spacing = 1.3
    intro.paragraph_format.space_after = Pt(8)
    r = intro.add_run(
        "Phần này trình bày các biểu đồ Use Case chi tiết cho từng nhóm chức năng chính của hệ thống Unilish. "
        "Mỗi biểu đồ mô tả tác nhân thực hiện chức năng và các thao tác nghiệp vụ bên trong phạm vi use case tương ứng."
    )
    set_font(r, 13)

    for idx, ((title, actor, _), img) in enumerate(zip(USECASES, image_paths), 2):
        if idx != 2:
            doc.add_section(WD_SECTION.NEW_PAGE)
        add_heading(doc, f"2.2.{idx-1}. Use case: {title}")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(str(img), width=Cm(14.5))
        add_caption(doc, f"Hình 2.{idx}. Use case Diagram {title}")

    doc.save(OUT)


def main():
    image_paths = []
    for i, (title, actor, items) in enumerate(USECASES, 1):
        svg = make_svg(i, title, actor, items)
        image_paths.append(convert_svg_to_png(svg))
    build_doc(image_paths)
    print(OUT)


if __name__ == "__main__":
    main()
