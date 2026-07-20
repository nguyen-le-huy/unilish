from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = Path("/Users/nguyenlehuy/Downloads/DATN2026-D2_CNTT_CNPM_0191966_NguyenLeHuy_BaoCao.docx")
OUT = Path("/Users/nguyenlehuy/Downloads/DATN2026-D2_CNTT_CNPM_0191966_NguyenLeHuy_BaoCao_SuaUseCaseTrang32.docx")


FIELDS = [
    "Use Case ID",
    "Use Case name",
    "Actor",
    "Priority",
    "Brief Description",
    "Pre-condition",
    "Trigger",
    "Basic Flows",
    "Alternative Flows",
    "Exception Flows",
    "Post-conditions",
    "Business Rules",
    "Non-Functional Requirements",
    "Special Requirements",
]


USE_CASES = [
    {
        "id": "UC-01",
        "name": "Đăng ký",
        "actor": "Learner",
        "priority": "Must",
        "brief": "Cho phép người học tạo tài khoản mới trên hệ thống Unilish để bắt đầu sử dụng các chức năng học ngoại ngữ.",
        "pre": "Người học chưa có tài khoản trong hệ thống.\nEmail đăng ký chưa tồn tại.",
        "trigger": "Người học chọn chức năng Đăng ký.",
        "basic": [
            "Người học mở màn hình đăng ký.",
            "Người học nhập họ tên, email và mật khẩu.",
            "Hệ thống kiểm tra tính hợp lệ của thông tin đăng ký.",
            "Hệ thống gửi mã OTP xác thực email.",
            "Người học nhập mã OTP.",
            "Hệ thống xác thực OTP và tạo tài khoản mới.",
            "Hệ thống chuyển người học đến bước đăng nhập hoặc thiết lập hồ sơ học tập.",
        ],
        "alt": "Người học có thể đăng ký/đăng nhập bằng tài khoản Google nếu tài khoản chưa tồn tại.",
        "ex": "Email đã được sử dụng.\nMật khẩu không đáp ứng yêu cầu bảo mật.\nMã OTP sai hoặc hết hạn.",
        "post": "Tài khoản người học được tạo và lưu trong hệ thống.",
        "rules": "Email không được trùng.\nMật khẩu phải đáp ứng yêu cầu bảo mật.\nTài khoản cần được xác thực trước khi sử dụng đầy đủ chức năng.",
        "nfr": "Thời gian phản hồi nhanh.\nThông tin mật khẩu phải được mã hóa.\nDữ liệu cá nhân được bảo vệ.",
        "special": "Hỗ trợ OTP qua email và Google OAuth.",
    },
    {
        "id": "UC-02",
        "name": "Đăng nhập",
        "actor": "Learner, Admin",
        "priority": "Must",
        "brief": "Cho phép người dùng truy cập hệ thống bằng tài khoản hợp lệ và được điều hướng theo đúng vai trò.",
        "pre": "Người dùng đã có tài khoản hợp lệ.",
        "trigger": "Người dùng chọn chức năng Đăng nhập.",
        "basic": [
            "Người dùng nhập email và mật khẩu.",
            "Hệ thống kiểm tra thông tin xác thực.",
            "Hệ thống xác định vai trò của người dùng.",
            "Hệ thống tạo phiên đăng nhập.",
            "Learner được chuyển đến giao diện học tập; Admin được chuyển đến trang quản trị.",
        ],
        "alt": "Người dùng có thể đăng nhập bằng Google nếu tài khoản đã liên kết.",
        "ex": "Sai email hoặc mật khẩu.\nTài khoản chưa xác thực hoặc bị khóa.\nToken đăng nhập không hợp lệ.",
        "post": "Người dùng đăng nhập thành công và có phiên làm việc hợp lệ.",
        "rules": "Mỗi vai trò chỉ được truy cập các chức năng được phân quyền.\nPhiên đăng nhập phải hết hạn sau một khoảng thời gian nhất định.",
        "nfr": "Bảo mật phiên đăng nhập.\nPhản hồi trong thời gian ngắn.\nKhông hiển thị thông tin lỗi nhạy cảm.",
        "special": "Sử dụng JWT hoặc cơ chế xác thực tương đương.",
    },
    {
        "id": "UC-03",
        "name": "Làm bài kiểm tra đầu vào",
        "actor": "Learner",
        "priority": "Must",
        "brief": "Đánh giá trình độ ban đầu của người học để hệ thống xác định cấp độ và đề xuất khóa học phù hợp.",
        "pre": "Người học đã đăng nhập.\nHệ thống có bài kiểm tra đầu vào đang hoạt động.",
        "trigger": "Người học chọn chức năng làm bài kiểm tra đầu vào.",
        "basic": [
            "Hệ thống hiển thị bài kiểm tra đầu vào.",
            "Người học trả lời các câu hỏi.",
            "Hệ thống lưu tạm câu trả lời trong quá trình làm bài.",
            "Người học nộp bài.",
            "Hệ thống chấm điểm và xác định trình độ.",
            "Hệ thống lưu kết quả vào hồ sơ học tập.",
            "Hệ thống sử dụng kết quả để gợi ý khóa học phù hợp.",
        ],
        "alt": "Người học có thể chọn trình độ thủ công nếu bỏ qua bài kiểm tra.",
        "ex": "Bài kiểm tra hết thời gian.\nMất kết nối khi nộp bài.\nBài kiểm tra chưa có đủ câu hỏi.",
        "post": "Kết quả kiểm tra và trình độ người học được lưu lại.",
        "rules": "Bài kiểm tra phải có thang điểm và tiêu chí phân loại trình độ.\nMỗi câu trả lời chỉ được tính điểm một lần.",
        "nfr": "Chấm điểm chính xác.\nLưu tiến trình ổn định.\nGiao diện làm bài dễ sử dụng.",
        "special": "Có thể chia theo kỹ năng nghe, đọc hoặc ngữ pháp.",
    },
    {
        "id": "UC-04",
        "name": "Luyện kỹ năng giao tiếp với AI",
        "actor": "Learner",
        "priority": "Should",
        "brief": "Cho phép người học luyện hội thoại với AI theo chủ đề nhằm cải thiện phản xạ giao tiếp ngoại ngữ.",
        "pre": "Người học đã đăng nhập.\nDịch vụ AI sẵn sàng hoạt động.",
        "trigger": "Người học chọn chức năng luyện giao tiếp với AI.",
        "basic": [
            "Người học chọn chủ đề hoặc tình huống giao tiếp.",
            "Hệ thống khởi tạo phiên hội thoại với AI.",
            "Người học nhập câu trả lời bằng văn bản hoặc giọng nói.",
            "AI phản hồi theo ngữ cảnh hội thoại.",
            "Hệ thống đưa ra gợi ý cải thiện nếu có.",
            "Hệ thống lưu lịch sử hoặc kết quả luyện tập.",
        ],
        "alt": "Người học có thể chuyển giữa luyện bằng văn bản và luyện bằng giọng nói.",
        "ex": "Dịch vụ AI không phản hồi.\nKhông thu được âm thanh.\nKết nối mạng không ổn định.",
        "post": "Phiên luyện giao tiếp được hoàn thành hoặc lưu lại.",
        "rules": "Nội dung phản hồi phải phù hợp với trình độ và chủ đề người học chọn.",
        "nfr": "Phản hồi AI nhanh, rõ ràng và dễ hiểu.\nDữ liệu hội thoại cần được bảo vệ.",
        "special": "Có thể tích hợp xử lý giọng nói và đánh giá câu trả lời.",
    },
    {
        "id": "UC-05",
        "name": "Học theo khóa học được đề xuất",
        "actor": "Learner",
        "priority": "Must",
        "brief": "Cho phép người học xem, đăng ký và học các khóa học được đề xuất dựa trên mục tiêu, ngôn ngữ và trình độ.",
        "pre": "Người học đã đăng nhập.\nHồ sơ học tập đã có thông tin ngôn ngữ, mục tiêu hoặc trình độ.",
        "trigger": "Người học mở danh sách khóa học được đề xuất.",
        "basic": [
            "Hệ thống hiển thị danh sách khóa học phù hợp.",
            "Người học chọn khóa học.",
            "Hệ thống hiển thị thông tin khóa học và lộ trình học.",
            "Người học đăng ký khóa học.",
            "Người học chọn unit hoặc lesson để học.",
            "Hệ thống hiển thị nội dung bài học và bài luyện tập.",
            "Hệ thống cập nhật tiến độ học tập sau khi người học hoàn thành bài học.",
        ],
        "alt": "Người học có thể tìm kiếm và học khóa học ngoài danh sách đề xuất.",
        "ex": "Khóa học chưa được xuất bản.\nLesson không có nội dung.\nKhông lưu được tiến độ học tập.",
        "post": "Khóa học được thêm vào lộ trình cá nhân và tiến độ được cập nhật.",
        "rules": "Chỉ hiển thị khóa học đã được xuất bản.\nLộ trình học phải theo thứ tự unit và lesson.",
        "nfr": "Nội dung bài học tải nhanh.\nGiao diện học tập rõ ràng, dễ theo dõi.",
        "special": "Hỗ trợ checkpoint để tiếp tục học từ vị trí gần nhất.",
    },
    {
        "id": "UC-06",
        "name": "Học với video YouTube",
        "actor": "Learner",
        "priority": "Could",
        "brief": "Cho phép người học học ngoại ngữ thông qua video YouTube được gắn với nội dung bài học.",
        "pre": "Người học đã đăng nhập.\nBài học có liên kết hoặc nội dung video hợp lệ.",
        "trigger": "Người học mở bài học có video YouTube.",
        "basic": [
            "Hệ thống hiển thị video trong giao diện bài học.",
            "Người học xem video.",
            "Người học thực hiện bài luyện tập liên quan nếu có.",
            "Hệ thống ghi nhận trạng thái hoàn thành nội dung video.",
            "Hệ thống cập nhật tiến độ bài học.",
        ],
        "alt": "Người học có thể mở video trên YouTube nếu trình nhúng không khả dụng.",
        "ex": "Video bị xóa hoặc không cho phép nhúng.\nMất kết nối mạng.\nĐường dẫn video không hợp lệ.",
        "post": "Tiến độ học với video được ghi nhận.",
        "rules": "Video phải phù hợp với nội dung bài học và trình độ người học.",
        "nfr": "Video hiển thị ổn định trên trình duyệt.\nKhông làm gián đoạn trải nghiệm học.",
        "special": "Lưu trữ URL video và mô tả nội dung trong dữ liệu lesson.",
    },
    {
        "id": "UC-07",
        "name": "Quản lý profile",
        "actor": "Learner",
        "priority": "Must",
        "brief": "Cho phép người học xem và cập nhật thông tin cá nhân, ngôn ngữ học, mục tiêu học tập và trình độ.",
        "pre": "Người học đã đăng nhập.",
        "trigger": "Người học mở trang profile.",
        "basic": [
            "Hệ thống hiển thị thông tin hồ sơ cá nhân.",
            "Người học chỉnh sửa thông tin cần cập nhật.",
            "Người học thay đổi ngôn ngữ, mục tiêu hoặc trình độ nếu cần.",
            "Hệ thống kiểm tra dữ liệu đầu vào.",
            "Hệ thống lưu thay đổi.",
            "Hệ thống hiển thị thông tin đã cập nhật.",
        ],
        "alt": "Người học chỉ xem thông tin mà không chỉnh sửa.",
        "ex": "Dữ liệu không hợp lệ.\nKhông lưu được thông tin.\nNgười học thay đổi email nhưng chưa xác thực.",
        "post": "Thông tin hồ sơ người học được cập nhật.",
        "rules": "Một số thông tin nhạy cảm cần xác thực khi thay đổi.\nHồ sơ học tập là cơ sở cho đề xuất khóa học.",
        "nfr": "Dữ liệu cá nhân được bảo mật.\nGiao diện chỉnh sửa đơn giản, dễ sử dụng.",
        "special": "Có thể hỗ trợ avatar, trình độ hiện tại và mục tiêu học tập.",
    },
    {
        "id": "UC-08",
        "name": "Luyện đề IELTS",
        "actor": "Learner",
        "priority": "Should",
        "brief": "Cho phép người học luyện đề IELTS theo từng kỹ năng như Listening, Reading, Writing và Speaking.",
        "pre": "Người học đã đăng nhập.\nHệ thống có đề luyện IELTS đã được xuất bản.",
        "trigger": "Người học chọn chức năng luyện đề IELTS.",
        "basic": [
            "Hệ thống hiển thị danh sách đề IELTS.",
            "Người học chọn đề và kỹ năng muốn luyện.",
            "Hệ thống hiển thị nội dung đề.",
            "Người học làm bài và nộp bài.",
            "Hệ thống chấm điểm hoặc lưu bài làm.",
            "Hệ thống hiển thị kết quả và lịch sử luyện tập.",
        ],
        "alt": "Với Writing hoặc Speaking, hệ thống có thể lưu bài làm để chấm hoặc phản hồi sau.",
        "ex": "Đề chưa đủ dữ liệu.\nLỗi âm thanh trong bài Listening.\nHết thời gian làm bài.",
        "post": "Kết quả luyện đề được lưu vào lịch sử học tập.",
        "rules": "Chỉ hiển thị đề đã được xuất bản.\nĐiểm số phải tính theo cấu trúc từng kỹ năng.",
        "nfr": "Đồng hồ làm bài và lưu câu trả lời hoạt động ổn định.",
        "special": "Có thể cần audio cho Listening và ghi âm cho Speaking.",
    },
    {
        "id": "UC-09",
        "name": "Quản lý ngôn ngữ",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên tạo, cập nhật và quản lý danh sách ngôn ngữ học được hỗ trợ trên hệ thống.",
        "pre": "Admin đã đăng nhập và có quyền quản trị.",
        "trigger": "Admin mở chức năng quản lý ngôn ngữ.",
        "basic": [
            "Hệ thống hiển thị danh sách ngôn ngữ.",
            "Admin thêm mới hoặc chỉnh sửa thông tin ngôn ngữ.",
            "Admin thiết lập trạng thái hoạt động của ngôn ngữ.",
            "Hệ thống kiểm tra dữ liệu.",
            "Hệ thống lưu thay đổi.",
        ],
        "alt": "Admin có thể tạm ẩn ngôn ngữ thay vì xóa khỏi hệ thống.",
        "ex": "Tên ngôn ngữ bị trùng.\nKhông thể xóa ngôn ngữ đang được dùng trong khóa học.",
        "post": "Danh sách ngôn ngữ được cập nhật.",
        "rules": "Ngôn ngữ đang có khóa học hoặc người học liên kết không được xóa trực tiếp.",
        "nfr": "Thao tác quản trị phải phản hồi nhanh và có thông báo kết quả.",
        "special": "Có thể lưu mã ngôn ngữ, tên hiển thị và trạng thái hoạt động.",
    },
    {
        "id": "UC-10",
        "name": "Quản lý mục tiêu",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên quản lý các mục tiêu học tập dùng trong quá trình onboarding và đề xuất khóa học.",
        "pre": "Admin đã đăng nhập và có quyền quản trị.",
        "trigger": "Admin mở chức năng quản lý mục tiêu học tập.",
        "basic": [
            "Hệ thống hiển thị danh sách mục tiêu học tập.",
            "Admin thêm, sửa hoặc thay đổi trạng thái mục tiêu.",
            "Admin liên kết mục tiêu với ngôn ngữ hoặc nhóm khóa học nếu cần.",
            "Hệ thống kiểm tra dữ liệu.",
            "Hệ thống lưu thay đổi.",
        ],
        "alt": "Admin có thể tạm ngừng sử dụng một mục tiêu mà không xóa dữ liệu.",
        "ex": "Tên mục tiêu bị trùng.\nMục tiêu đang được người học sử dụng nên không thể xóa.",
        "post": "Danh sách mục tiêu học tập được cập nhật.",
        "rules": "Mục tiêu học tập là dữ liệu đầu vào cho hồ sơ người học và đề xuất khóa học.",
        "nfr": "Dữ liệu quản trị phải nhất quán và dễ tra cứu.",
        "special": "Có thể phân loại mục tiêu theo nhu cầu như giao tiếp, thi chứng chỉ hoặc công việc.",
    },
    {
        "id": "UC-11",
        "name": "Quản lý nội dung học",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên quản lý khóa học, unit, lesson và nội dung học tập trong hệ thống Unilish.",
        "pre": "Admin đã đăng nhập.\nHệ thống đã có ngôn ngữ và mục tiêu học tập cần thiết.",
        "trigger": "Admin mở khu vực quản lý nội dung học.",
        "basic": [
            "Hệ thống hiển thị danh sách khóa học, unit và lesson.",
            "Admin tạo mới hoặc chỉnh sửa khóa học.",
            "Admin thêm unit và lesson vào khóa học.",
            "Admin cập nhật nội dung bài học, bài luyện tập và tài nguyên liên quan.",
            "Admin lưu bản nháp.",
            "Admin gửi nội dung để duyệt hoặc xuất bản.",
        ],
        "alt": "Admin có thể dùng chức năng AI để hỗ trợ sinh nội dung hoặc câu hỏi.",
        "ex": "Thiếu thông tin bắt buộc.\nNội dung lesson chưa hợp lệ.\nKhông thể xuất bản khóa học chưa có lesson.",
        "post": "Nội dung học được lưu, cập nhật hoặc chuyển sang trạng thái xuất bản.",
        "rules": "Khóa học phải có cấu trúc unit và lesson rõ ràng.\nChỉ nội dung đã xuất bản mới hiển thị cho Learner.",
        "nfr": "Giao diện quản trị nội dung phải ổn định khi dữ liệu lớn.",
        "special": "Hỗ trợ nhiều dạng nội dung như vocabulary, grammar, reading, listening, writing và speaking.",
    },
    {
        "id": "UC-12",
        "name": "Quản lý user",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên xem, tìm kiếm và quản lý tài khoản người dùng trong hệ thống.",
        "pre": "Admin đã đăng nhập và có quyền quản lý người dùng.",
        "trigger": "Admin mở chức năng quản lý user.",
        "basic": [
            "Hệ thống hiển thị danh sách người dùng.",
            "Admin tìm kiếm hoặc lọc người dùng.",
            "Admin xem chi tiết thông tin tài khoản.",
            "Admin cập nhật vai trò, cấp độ hoặc trạng thái tài khoản.",
            "Hệ thống lưu thay đổi và ghi nhận kết quả.",
        ],
        "alt": "Admin có thể vô hiệu hóa tài khoản thay vì xóa vĩnh viễn.",
        "ex": "Không tìm thấy người dùng.\nAdmin không đủ quyền thay đổi vai trò.\nKhông thể xóa tài khoản đang có dữ liệu liên quan.",
        "post": "Thông tin tài khoản người dùng được cập nhật.",
        "rules": "Không cho phép Admin tự xóa hoặc hạ quyền tài khoản quản trị quan trọng nếu hệ thống không cho phép.",
        "nfr": "Danh sách người dùng cần hỗ trợ tìm kiếm và phân trang.",
        "special": "Có thể quản lý vai trò Learner, Admin hoặc các vai trò nội dung.",
    },
    {
        "id": "UC-13",
        "name": "Quản lý ngân hàng câu hỏi",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên tạo, chỉnh sửa, phân loại và duyệt các câu hỏi dùng trong lesson, placement test và IELTS.",
        "pre": "Admin đã đăng nhập.\nHệ thống có cấu hình ngôn ngữ, cấp độ hoặc bài học liên quan.",
        "trigger": "Admin mở chức năng quản lý ngân hàng câu hỏi.",
        "basic": [
            "Hệ thống hiển thị danh sách câu hỏi.",
            "Admin tạo mới hoặc chỉnh sửa câu hỏi.",
            "Admin nhập đáp án, giải thích, cấp độ và kỹ năng liên quan.",
            "Admin liên kết câu hỏi với lesson, placement test hoặc đề IELTS.",
            "Admin lưu hoặc gửi duyệt câu hỏi.",
            "Hệ thống cập nhật trạng thái câu hỏi.",
        ],
        "alt": "Admin có thể sử dụng AI để sinh câu hỏi nháp rồi chỉnh sửa trước khi duyệt.",
        "ex": "Thiếu đáp án đúng.\nCâu hỏi trùng nội dung.\nKhông thể xuất bản câu hỏi chưa được duyệt.",
        "post": "Ngân hàng câu hỏi được cập nhật và sẵn sàng sử dụng cho các bài học hoặc bài kiểm tra.",
        "rules": "Câu hỏi phải có đáp án đúng và metadata phù hợp.\nChỉ câu hỏi đã duyệt mới được dùng trong bài chính thức.",
        "nfr": "Hỗ trợ tìm kiếm, lọc và xử lý số lượng câu hỏi lớn.",
        "special": "Có thể hỗ trợ nhiều dạng câu hỏi như trắc nghiệm, điền từ, nghe hiểu hoặc tự luận.",
    },
    {
        "id": "UC-14",
        "name": "Quản lý bài kiểm tra đầu vào",
        "actor": "Admin",
        "priority": "Must",
        "brief": "Cho phép quản trị viên tạo và cấu hình bài kiểm tra đầu vào để đánh giá trình độ người học.",
        "pre": "Admin đã đăng nhập.\nNgân hàng câu hỏi có câu hỏi phù hợp.",
        "trigger": "Admin mở chức năng quản lý bài kiểm tra đầu vào.",
        "basic": [
            "Hệ thống hiển thị danh sách bài kiểm tra đầu vào.",
            "Admin tạo mới hoặc chỉnh sửa bài kiểm tra.",
            "Admin chọn kỹ năng, cấp độ, số lượng câu hỏi và thời gian làm bài.",
            "Admin liên kết câu hỏi từ ngân hàng câu hỏi.",
            "Admin cấu hình thang điểm và quy tắc phân loại trình độ.",
            "Admin xuất bản bài kiểm tra.",
        ],
        "alt": "Admin có thể tạo phiên bản nháp để kiểm tra trước khi xuất bản.",
        "ex": "Bài kiểm tra chưa có đủ câu hỏi.\nThang điểm không hợp lệ.\nKhông thể xuất bản do thiếu cấu hình bắt buộc.",
        "post": "Bài kiểm tra đầu vào được lưu và có thể sử dụng cho Learner.",
        "rules": "Bài kiểm tra phải có đủ câu hỏi và quy tắc tính điểm trước khi xuất bản.",
        "nfr": "Cấu hình bài kiểm tra phải rõ ràng và dễ chỉnh sửa.",
        "special": "Có thể hỗ trợ nhiều phiên bản bài kiểm tra theo ngôn ngữ hoặc trình độ.",
    },
    {
        "id": "UC-15",
        "name": "Quản lý đề IELTS",
        "actor": "Admin",
        "priority": "Should",
        "brief": "Cho phép quản trị viên tạo, chỉnh sửa và xuất bản các đề luyện IELTS theo từng kỹ năng.",
        "pre": "Admin đã đăng nhập.\nHệ thống có dữ liệu câu hỏi, audio hoặc nội dung đề phù hợp.",
        "trigger": "Admin mở chức năng quản lý đề IELTS.",
        "basic": [
            "Hệ thống hiển thị danh sách đề IELTS.",
            "Admin tạo mới hoặc chỉnh sửa đề.",
            "Admin cấu hình kỹ năng Listening, Reading, Writing hoặc Speaking.",
            "Admin thêm câu hỏi, đoạn văn, audio hoặc yêu cầu bài làm.",
            "Admin kiểm tra nội dung đề.",
            "Admin xuất bản đề luyện.",
        ],
        "alt": "Admin có thể lưu bản nháp hoặc tạo phiên bản mới từ đề đã có.",
        "ex": "Thiếu audio cho Listening.\nThiếu nội dung câu hỏi.\nKhông thể xuất bản đề chưa hoàn chỉnh.",
        "post": "Đề IELTS được lưu và có thể hiển thị cho người học luyện tập.",
        "rules": "Đề phải có cấu trúc đúng theo kỹ năng.\nChỉ đề đã xuất bản mới hiển thị cho Learner.",
        "nfr": "Âm thanh và nội dung đề phải tải ổn định.",
        "special": "Có thể hỗ trợ phiên bản đề và trạng thái nháp/xuất bản/lưu trữ.",
    },
]


def set_run_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_para_font(paragraph, size=13, bold=False, italic=False):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def set_cell_text(cell, value, size=11):
    cell.text = ""
    lines = str(value).split("\n") if value else [""]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(line)
        set_run_font(run, size=size)


def set_numbered_cell(cell, items, size=11):
    cell.text = ""
    for i, item in enumerate(items, 1):
        p = cell.paragraphs[0] if i == 1 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(f"{i}. {item}")
        set_run_font(run, size=size)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            tcW = cell._tc.get_or_add_tcPr().tcW
            tcW.type = "dxa"
            tcW.w = int(width * 567)


def add_spec_table(doc, uc, index):
    title = doc.add_paragraph()
    title.style = "List Paragraph"
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.first_line_indent = None
    run = title.add_run(f"2.1.{index}. Use case: {uc['name']}")
    set_run_font(run, 13, bold=True)

    values = {
        "Use Case ID": uc["id"],
        "Use Case name": uc["name"],
        "Actor": uc["actor"],
        "Priority": uc["priority"],
        "Brief Description": uc["brief"],
        "Pre-condition": uc["pre"],
        "Trigger": uc["trigger"],
        "Basic Flows": uc["basic"],
        "Alternative Flows": uc["alt"],
        "Exception Flows": uc["ex"],
        "Post-conditions": uc["post"],
        "Business Rules": uc["rules"],
        "Non-Functional Requirements": uc["nfr"],
        "Special Requirements": uc["special"],
    }

    table = doc.add_table(rows=len(FIELDS), cols=2)
    set_table_width(table, [4.1, 10.6])

    for row, field in zip(table.rows, FIELDS):
        left, right = row.cells
        set_cell_text(left, field, 11)
        for p in left.paragraphs:
            for r in p.runs:
                r.bold = True if field in ("Use Case ID", "Use Case name", "Actor", "Priority") else False
        if field == "Basic Flows":
            set_numbered_cell(right, values[field], 11)
        else:
            set_cell_text(right, values[field], 11)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)
            set_cell_margins(cell)

    caption = doc.add_paragraph()
    caption.style = "Caption"
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run(f"Bảng 2.{index}. Mô tả UC {uc['name']}")
    set_run_font(r, 13, italic=True)

    return [title._element, table._element, caption._element]


def body_text(element):
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def main():
    doc = Document(SRC)

    for style_name in ("Normal", "Body Text", "List Paragraph", "Caption"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            style.font.size = Pt(13)

    body = doc.element.body
    start = None
    anchor = None
    for i, child in enumerate(body):
        txt = body_text(child)
        if start is None and txt.startswith("Use case: Khám phá dungeon"):
            start = i
        if txt == "Thiết kế hệ thống":
            anchor = child
            break
    if start is None or anchor is None:
        raise RuntimeError("Không tìm thấy vùng use case cũ hoặc điểm neo Thiết kế hệ thống.")

    while body[start] is not anchor:
        body.remove(body[start])

    new_elements = []
    for idx, uc in enumerate(USE_CASES, 1):
        new_elements.extend(add_spec_table(doc, uc, idx))

    for element in new_elements:
        body.remove(element)
        anchor.addprevious(element)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
